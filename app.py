import joblib
from flask import Flask, request, render_template, jsonify, send_file, redirect, url_for, flash, session, abort
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
import openai
from docx import Document
from datetime import datetime
from docx.shared import Pt
import os
import logging
from functools import wraps
import json
import re
import fitz  # PyMuPDF

# --- New imports for ML model ---
from sklearn.datasets import load_files
from sklearn.pipeline import Pipeline
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key-here')  # Use environment variable if available
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///cv_maker.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 *1024 # 16MB max file size

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

db = SQLAlchemy(app)

# ------------------------------
# Login required decorator
# ------------------------------
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please log in to access this page.')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# ------------------------------
# Database Models
# ------------------------------
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(128))
    cvs = db.relationship('CV', backref='user', lazy=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class CV(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

# Create database tables
with app.app_context():
    try:
        db.create_all()
        logger.info("Database tables created successfully")
    except Exception as e:
        logger.error(f"Error creating database tables: {e}")

# ------------------------------
# Error Handlers
# ------------------------------
@app.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('500.html'), 500

@app.errorhandler(413)
def too_large(error):
    return render_template('413.html'), 413

# ------------------------------
# OpenAI API Key (replace with your key)
# ------------------------------
openai.api_key = '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'
# ------------------------------
# CV Chat & Generation Steps (existing functionality)
# ------------------------------
CV_STEPS = {
    'personal_info': {
        'name': 'What is your full name?',
        'email': 'What is your email address?',
        'phone': 'What is your phone number?',
        'location': 'Where are you located?',
        'summary': 'Brief professional summary (2-3 sentences)?',
        'languages': 'Languages you speak? (e.g., English - Native)',
        'hobbies': 'Hobbies and interests?'
    },
    'experience': {
        'experience': 'List your work experiences. For each: Company, Title, Start Date, End Date, Description, Achievements.'
    },
    'education': {
        'education': 'List your education. For each: Institution, Degree, Field, Start Date, End Date.'
    },
    'skills': {
        'technical': 'Technical skills? (e.g., Python, JavaScript)',
        'soft': 'Soft skills? (e.g., Leadership, Communication)',
        'certifications': 'Certifications (if any)?'
    }
}

user_data = {}
step_keys = list(CV_STEPS.keys())
step_index = 0
question_index = 0

def get_next_question():
    current_step = step_keys[step_index]
    current_questions = list(CV_STEPS[current_step].keys())
    return CV_STEPS[current_step][current_questions[question_index]]

# ------------------------------
# Routes for Registration, Login, Logout, Profile, etc. (existing functionality)
# ------------------------------
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        try:
            username = request.form.get('username')
            email = request.form.get('email')
            password = request.form.get('password')

            if not username or not email or not password:
                flash('All fields are required')
                return redirect(url_for('register'))

            if User.query.filter_by(username=username).first():
                flash('Username already exists')
                return redirect(url_for('register'))

            if User.query.filter_by(email=email).first():
                flash('Email already registered')
                return redirect(url_for('register'))

            user = User(username=username, email=email)
            user.set_password(password)
            db.session.add(user)
            db.session.commit()

            flash('Registration successful! Please login.')
            return redirect(url_for('login'))
        except Exception as e:
            db.session.rollback()
            logger.error(f"Registration error: {e}")
            flash('An error occurred during registration. Please try again.')
            return redirect(url_for('register'))

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        try:
            username = request.form.get('username')
            password = request.form.get('password')
            
            if not username or not password:
                flash('Please enter both username and password')
                return redirect(url_for('login'))
                
            user = User.query.filter_by(username=username).first()

            if user and user.check_password(password):
                session['user_id'] = user.id
                session['username'] = user.username
                flash('Logged in successfully!')
                return redirect(url_for('home'))
            flash('Invalid username or password')
        except Exception as e:
            logger.error(f"Login error: {e}")
            flash('An error occurred during login. Please try again.')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('Logged out successfully!')
    return redirect(url_for('login'))

@app.route('/profile')
@login_required
def profile():
    try:
        user = User.query.get(session['user_id'])
        cvs = CV.query.filter_by(user_id=user.id).order_by(CV.created_at.desc()).all()
        return render_template('profile.html', user=user, cvs=cvs)
    except Exception as e:
        logger.error(f"Profile error: {e}")
        flash('An error occurred while loading your profile.')
        return redirect(url_for('home'))

@app.route('/edit_profile', methods=['POST'])
@login_required
def edit_profile():
    try:
        user = User.query.get(session['user_id'])
        username = request.form.get('username')
        email = request.form.get('email')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')

        if username != user.username:
            existing_user = User.query.filter_by(username=username).first()
            if existing_user:
                flash('Username already exists')
                return redirect(url_for('profile'))

        if email != user.email:
            existing_user = User.query.filter_by(email=email).first()
            if existing_user:
                flash('Email already registered')
                return redirect(url_for('profile'))

        user.username = username
        user.email = email

        if new_password:
            if new_password != confirm_password:
                flash('New passwords do not match')
                return redirect(url_for('profile'))
            user.set_password(new_password)

        db.session.commit()
        session['username'] = username
        flash('Profile updated successfully!')
    except Exception as e:
        db.session.rollback()
        logger.error(f"Profile update error: {e}")
        flash('An error occurred while updating your profile.')
    return redirect(url_for('profile'))

@app.route('/delete_cv/<int:cv_id>')
@login_required
def delete_cv(cv_id):
    try:
        cv = CV.query.get_or_404(cv_id)
        if cv.user_id != session['user_id']:
            abort(403)
        if os.path.exists(cv.filename):
            os.remove(cv.filename)
        db.session.delete(cv)
        db.session.commit()
        flash('CV deleted successfully!')
    except Exception as e:
        db.session.rollback()
        logger.error(f"Delete CV error: {e}")
        flash('Error deleting CV')
    return redirect(url_for('profile'))

@app.route('/analyze_cv/<int:cv_id>')
@login_required
def analyze_cv(cv_id):
    try:
        cv = CV.query.get_or_404(cv_id)
        if cv.user_id != session['user_id']:
            abort(403)
        doc = Document(cv.filename)
        cv_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        analysis_prompt = f"""
        Analyze this CV and provide a detailed assessment. Format the response in plain text with no Markdown formatting.
        Include sections: Key skills, years of experience & achievements, important keywords, overall score out of 100, strengths & improvement areas, recommendations.
        CV Content:
        {cv_text}
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": analysis_prompt}],
            temperature=0.7
        )
        analysis = response['choices'][0]['message']['content']
        return render_template('cv_analysis.html', cv=cv, analysis=analysis)
    except Exception as e:
        logger.error(f"CV analysis error: {e}")
        flash('An error occurred while analyzing your CV.')
        return redirect(url_for('profile'))

@app.route('/analyze_upload', methods=['GET', 'POST'])
@login_required
def analyze_upload():
    if request.method == 'POST':
        try:
            if 'cv_file' not in request.files:
                flash('No file uploaded')
                return redirect(request.url)
            file = request.files['cv_file']
            if file.filename == '':
                flash('No file selected')
                return redirect(request.url)
            if not file:
                flash('Invalid file')
                return redirect(request.url)
            temp_dir = 'static/temp'
            os.makedirs(temp_dir, exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            temp_filename = f"{temp_dir}/temp_{session['user_id']}_{timestamp}{os.path.splitext(file.filename)[1]}"
            file.save(temp_filename)
            session['uploaded_cv'] = temp_filename
            filename = file.filename.lower()
            if filename.endswith('.docx'):
                doc = Document(temp_filename)
                cv_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            elif filename.endswith('.txt'):
                with open(temp_filename, 'r', encoding='utf-8') as f:
                    cv_text = f.read()
            elif filename.endswith('.pdf'):
                from PyPDF2 import PdfReader
                reader = PdfReader(temp_filename)
                cv_text = ""
                for page in reader.pages:
                    cv_text += page.extract_text() + "\n"
            else:
                flash('Unsupported file format. Please upload DOCX, PDF, or TXT files.')
                return redirect(request.url)
            analysis_prompt = f"""
            Analyze this CV and provide a detailed assessment in plain text (no Markdown formatting). 
            Sections: Key skills, experience & achievements, keywords, overall score, strengths & improvement areas, recommendations.
            CV Content:
            {cv_text}
            """
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": analysis_prompt}],
                temperature=0.7
            )
            analysis = response['choices'][0]['message']['content']
            return render_template('analyze_upload.html', analysis=analysis)
        except Exception as e:
            logger.error(f"CV upload analysis error: {e}")
            flash('An error occurred while analyzing your CV.')
            return redirect(request.url)
    return render_template('analyze_upload.html')

@app.route('/jobs')
@login_required
def jobs():
    try:
        user = User.query.get(session['user_id'])
        latest_cv = CV.query.filter_by(user_id=user.id).order_by(CV.created_at.desc()).first()
        if not latest_cv:
            flash('Please create or upload a CV to get job recommendations.')
            return redirect(url_for('home'))
        doc = Document(latest_cv.filename)
        cv_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        job_prompt = f"""
        Based on this CV, suggest 5 relevant job opportunities. For each, provide:
        Job title, Company name, Location, Experience level required, Brief description, Required skills, Application URL.
        CV Content:
        {cv_text}
        Format the response as a JSON array.
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": job_prompt}],
            temperature=0.7
        )
        jobs = response['choices'][0]['message']['content']
        try:
            jobs = json.loads(jobs)
        except json.JSONDecodeError:
            jobs = []
        job_title = request.args.get('job_title', '')
        location = request.args.get('location', '')
        experience_level = request.args.get('experience_level', '')
        if job_title or location or experience_level:
            filtered_jobs = []
            for job in jobs:
                if job_title and job_title.lower() not in job['title'].lower():
                    continue
                if location and location.lower() not in job['location'].lower():
                    continue
                if experience_level and experience_level.lower() != job['experience_level'].lower():
                    continue
                filtered_jobs.append(job)
            jobs = filtered_jobs
        return render_template('jobs.html', jobs=jobs)
    except Exception as e:
        logger.error(f"Job recommendations error: {e}")
        flash('An error occurred while getting job recommendations.')
        return redirect(url_for('home'))

@app.route('/')
def home():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    global step_index, question_index, user_data
    step_index = 0
    question_index = 0
    user_data = {}
    return render_template('chat.html', question=get_next_question())

@app.route('/chat', methods=['POST'])
def chat():
    global step_index, question_index, user_data
    answer = request.json.get('answer')
    current_step = step_keys[step_index]
    current_questions = list(CV_STEPS[current_step].keys())
    key = current_questions[question_index]
    user_data[key] = answer
    question_index += 1
    if question_index >= len(current_questions):
        step_index += 1
        question_index = 0
    if step_index >= len(step_keys):
        return jsonify({'done': True})
    return jsonify({'question': get_next_question()})

@app.route('/generate', methods=['GET'])
def generate_cv():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    cv_id = request.args.get('cv_id')
    if cv_id:
        cv = CV.query.get_or_404(cv_id)
        if cv.user_id != session['user_id']:
            flash('Unauthorized access')
            return redirect(url_for('profile'))
        return send_file(cv.filename, as_attachment=True)
    prompt = f"""
    You are a CV writing expert. Use the following data to generate a clean, professional CV.
    Data: {user_data}
    Output only the final CV content Please if the cv data is illogical please put in the cv that its illogical.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    final_cv_text = response['choices'][0]['message']['content']
    os.makedirs('static/cvs', exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    filename = f"static/cvs/cv_{session['user_id']}_{timestamp}.docx"
    doc = Document()
    doc.add_heading('Curriculum Vitae', 0)
    for line in final_cv_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.isupper() or line.endswith(":"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    doc.save(filename)
    cv = CV(filename=filename, user_id=session['user_id'])
    db.session.add(cv)
    db.session.commit()
    return send_file(filename, as_attachment=True)

@app.route('/fix_cv_issues/<int:cv_id>')
@login_required
def fix_cv_issues(cv_id):
    try:
        cv = CV.query.get_or_404(cv_id)
        if cv.user_id != session['user_id']:
            abort(403)
        doc = Document(cv.filename)
        cv_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        improvement_prompt = f"""
        You are a CV writing expert. Analyze this CV and create an improved version that addresses all issues.
        CV Content:
        {cv_text}
        Provide the improved CV in plain text.
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": improvement_prompt}],
            temperature=0.7
        )
        improved_cv_text = response['choices'][0]['message']['content']
        os.makedirs('static/cvs', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        improved_filename = f"static/cvs/cv_{session['user_id']}_{timestamp}_improved.docx"
        doc = Document()
        doc.add_heading('Curriculum Vitae', 0)
        for line in improved_cv_text.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
            elif line.isupper() or line.endswith(":"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        doc.save(improved_filename)
        improved_cv = CV(filename=improved_filename, user_id=session['user_id'])
        db.session.add(improved_cv)
        db.session.commit()
        flash('Your CV has been improved! Download the new version below.')
        return redirect(url_for('analyze_cv', cv_id=improved_cv.id))
    except Exception as e:
        logger.error(f"CV improvement error: {e}")
        flash('An error occurred while improving your CV.')
        return redirect(url_for('profile'))

@app.route('/fix_uploaded_cv')
@login_required
def fix_uploaded_cv():
    try:
        if 'uploaded_cv' not in session:
            flash('No CV file found. Please upload a CV first.')
            return redirect(url_for('analyze_upload'))
        file_path = session['uploaded_cv']
        if not os.path.exists(file_path):
            flash('CV file not found. Please upload again.')
            return redirect(url_for('analyze_upload'))
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            cv_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                cv_text = f.read()
        elif file_path.endswith('.pdf'):
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            cv_text = ""
            for page in reader.pages:
                cv_text += page.extract_text() + "\n"
        improvement_prompt = f"""
        You are a CV writing expert. Analyze this CV and create an improved version.
        CV Content:
        {cv_text}
        Provide the improved CV in plain text.
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": improvement_prompt}],
            temperature=0.7
        )
        improved_cv_text = response['choices'][0]['message']['content']
        os.makedirs('static/cvs', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        improved_filename = f"static/cvs/cv_{session['user_id']}_{timestamp}_improved.docx"
        doc = Document()
        doc.add_heading('Curriculum Vitae', 0)
        for line in improved_cv_text.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
            elif line.isupper() or line.endswith(":"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        doc.save(improved_filename)
        improved_cv = CV(filename=improved_filename, user_id=session['user_id'])
        db.session.add(improved_cv)
        db.session.commit()
        if os.path.exists(file_path):
            os.remove(file_path)
        session.pop('uploaded_cv', None)
        flash('Your CV has been improved! Download the new version below.')
        return redirect(url_for('analyze_cv', cv_id=improved_cv.id))
    except Exception as e:
        logger.error(f"CV improvement error: {e}")
        flash('An error occurred while improving your CV.')
        return redirect(url_for('analyze_upload'))


model = joblib.load("Models/Career Advisor/carrer_model.pkl")
label_encoder = joblib.load("Models/Career Advisor/career_label_encoder.pkl")

@app.route("/predict-career", methods=["GET", "POST"])
def predict_career():
    prediction = None

    if request.method == "POST":
        # Collect input values from the form
        form_data = {
            "GPA": float(request.form["GPA"]),
            "Extracurricular_Activities": int(request.form["Extracurricular_Activities"]),
            "Internships": int(request.form["Internships"]),
            "Projects": int(request.form["Projects"]),
            "Leadership_Positions": int(request.form["Leadership_Positions"]),
            "Field_Specific_Courses": int(request.form["Field_Specific_Courses"]),
            "Research_Experience": int(request.form["Research_Experience"]),
            "Coding_Skills": int(request.form["Coding_Skills"]),
            "Communication_Skills": int(request.form["Communication_Skills"]),
            "Problem_Solving_Skills": int(request.form["Problem_Solving_Skills"]),
            "Teamwork_Skills": int(request.form["Teamwork_Skills"]),
            "Analytical_Skills": int(request.form["Analytical_Skills"]),
            "Presentation_Skills": int(request.form["Presentation_Skills"]),
            "Networking_Skills": int(request.form["Networking_Skills"]),
            "Industry_Certifications": int(request.form["Industry_Certifications"]),
            "Field": request.form["Field"]
        }

        # Make sure the order of inputs matches training
        input_df = pd.DataFrame([form_data])
        y_pred_encoded = model.predict(input_df)[0]
        prediction = label_encoder.inverse_transform([y_pred_encoded])[0]

    return render_template("predict_career.html", prediction=prediction)


# -----------------------------
# Skills Matching Functions
# -----------------------------
def skill_tokenizer(text):
    return text.split(', ')

def clean_text(text):
    text = text.lower()
    text = re.sub(r'[^a-zA-Z0-9, ]', '', text)
    return text

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    return "\n".join(page.get_text() for page in doc)

def extract_skills(resume_text, known_skills):
    resume_text = resume_text.lower()
    return [skill for skill in known_skills if re.search(rf'\b{re.escape(skill)}\b', resume_text)]

@app.route('/skills-matching', methods=['GET', 'POST'])
@login_required
def skills_matching():
    prediction = None
    extracted_skills = []
    
    if request.method == 'POST':
        if 'resume_file' not in request.files:
            flash('No file uploaded')
            return redirect(request.url)
        
        file = request.files['resume_file']
        if file.filename == '':
            flash('No file selected')
            return redirect(request.url)
            
        if not file.filename.lower().endswith('.pdf'):
            flash('Please upload a PDF file')
            return redirect(request.url)
            
        try:
            # Create temp directory if it doesn't exist
            temp_dir = 'static/temp'
            os.makedirs(temp_dir, exist_ok=True)
            
            # Save uploaded file
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            pdf_path = f"{temp_dir}/resume_{session['user_id']}_{timestamp}.pdf"
            file.save(pdf_path)
            
            # Load model and label encoder
            skills_model_dir = "Models/Skills Matching"
            model = joblib.load(os.path.join(skills_model_dir, "career_classifier_advanced.pkl"))
            label_encoder = joblib.load(os.path.join(skills_model_dir, "label_encoder.pkl"))
            
            # Extract text and skills
            resume_text = extract_text_from_pdf(pdf_path)
            known_skills = list(model.named_steps['tfidf'].vocabulary_.keys())
            extracted_skills = extract_skills(resume_text, known_skills)
            
            if not extracted_skills:
                flash("No recognizable skills found in your resume.")
                return render_template('skills_matching.html', prediction=None, extracted_skills=[])
                
            cleaned_input = clean_text(", ".join(extracted_skills))
            prediction_encoded = model.predict([cleaned_input])
            prediction = label_encoder.inverse_transform(prediction_encoded)[0]
            
            # Clean up - Remove temporary file
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
                
        except Exception as e:
            logger.error(f"Skills matching error: {e}")
            flash(f'An error occurred: {str(e)}')
            return redirect(request.url)
            
    return render_template('skills_matching.html', prediction=prediction, extracted_skills=extracted_skills)


if __name__ == '__main__':
    app.run(debug=True,port=5001)
